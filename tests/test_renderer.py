"""Tests for output rendering."""

from docx import Document

from cvagent.renderer import render_docx, render_markdown, render_text


def test_render_markdown_creates_file(tmp_path):
	out = tmp_path / "letter.md"
	render_markdown("# Hello **World**", out)
	assert out.exists()
	assert out.read_text() == "# Hello **World**"


def test_render_markdown_preserves_bold_markers(tmp_path):
	out = tmp_path / "letter.md"
	render_markdown("I am **experienced** in Python.", out)
	assert "**experienced**" in out.read_text()


def test_render_text_creates_file(tmp_path):
	out = tmp_path / "letter.txt"
	render_text("Hello World", out)
	assert out.exists()


def test_render_text_strips_bold_markers(tmp_path):
	out = tmp_path / "letter.txt"
	render_text("Hello **World** and **Python**", out)
	content = out.read_text()
	assert "**" not in content
	assert "Hello World and Python" in content


def test_render_markdown_creates_parent_dirs(tmp_path):
	out = tmp_path / "nested" / "dir" / "letter.md"
	render_markdown("Hello", out)
	assert out.exists()


# ---------------------------------------------------------------------------
# render_docx tests
# ---------------------------------------------------------------------------


def _sample_profile() -> dict:
	return {
		"contact": {
			"name": "Jane Doe",
			"email": "jane@example.com",
			"phone": "+1 555 000 0000",
			"location": "Helsinki, Finland",
		},
		"skills": ["Python"],
		"work_history": [{"company": "Acme", "title": "Engineer", "start": "2021-01"}],
		"education": [{"degree": "M.Sc. Computer Science", "institution": "University of Helsinki", "year": 2021}],
	}


def _sample_job() -> dict:
	return {"company": "Futurice", "role": "Senior Python Engineer", "description": "A job."}


def _extract_all_text(doc) -> str:
	"""Pull all text out of paragraphs and table cells."""
	parts = [p.text for p in doc.paragraphs]
	for table in doc.tables:
		for row in table.rows:
			for cell in row.cells:
				parts.append(cell.text)
	return " ".join(parts)


def test_render_docx_creates_file(tmp_path):
	out = tmp_path / "letter.docx"
	render_docx("Hello world.", out, _sample_profile(), _sample_job())
	assert out.exists()
	assert out.suffix == ".docx"


def test_render_docx_contains_name(tmp_path):
	out = tmp_path / "letter.docx"
	render_docx("Hello world.", out, _sample_profile(), _sample_job())
	all_text = _extract_all_text(Document(out))
	assert "Jane Doe" in all_text


def test_render_docx_contains_company(tmp_path):
	out = tmp_path / "letter.docx"
	render_docx("Hello world.", out, _sample_profile(), _sample_job())
	all_text = _extract_all_text(Document(out))
	assert "Futurice" in all_text


def test_render_docx_contains_body_text(tmp_path):
	out = tmp_path / "letter.docx"
	render_docx("Dear Hiring Team, I am applying.", out, _sample_profile(), _sample_job())
	all_text = _extract_all_text(Document(out))
	assert "Dear Hiring Team" in all_text


def test_render_docx_strips_bold_markers(tmp_path):
	out = tmp_path / "letter.docx"
	render_docx("I am **experienced** in Python.", out, _sample_profile(), _sample_job())
	all_text = _extract_all_text(Document(out))
	assert "**" not in all_text
	assert "experienced" in all_text


def test_render_docx_creates_parent_dirs(tmp_path):
	out = tmp_path / "nested" / "dir" / "letter.docx"
	render_docx("Hello", out, _sample_profile(), _sample_job())
	assert out.exists()


def test_render_docx_has_two_tables(tmp_path):
	out = tmp_path / "letter.docx"
	render_docx("Hello", out, _sample_profile(), _sample_job())
	doc = Document(out)
	assert len(doc.tables) == 2


def test_render_docx_credential_fallback(tmp_path):
	profile = {
		"contact": {"name": "Bob Smith", "email": "b@b.com", "phone": "123", "location": "London"},
		"skills": [],
		"work_history": [{"company": "Corp", "title": "Developer", "start": "2020-01"}],
	}
	out = tmp_path / "letter.docx"
	render_docx("Hello", out, profile, _sample_job())
	all_text = _extract_all_text(Document(out))
	assert "Developer" in all_text


def test_render_docx_missing_optional_fields(tmp_path):
	profile = {
		"contact": {"name": "No Phone"},
		"skills": [],
		"work_history": [],
	}
	out = tmp_path / "letter.docx"
	render_docx("Hello", out, profile, {"company": "Corp", "role": "Dev"})
	assert out.exists()
