"""Format and save the generated cover letter."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor

_HEADER_FILL = "2E4057"  # dark navy
_SIDEBAR_FILL = "F0F0F0"  # light grey
_SIDEBAR_WIDTH = Cm(3.4)
_CONTENT_WIDTH = Cm(13.6)
_MARGIN = Cm(1.8)
_FOOTER_TEXT = (
	"Generated with CVAgent — an open-source cover letter automation tool developed by Eero Isola. "
	"The project itself is a demonstration of practical software engineering skills."
)


def render_markdown(text: str, output_path: Path) -> None:
	"""Write the cover letter as a .md file."""
	output_path.parent.mkdir(parents=True, exist_ok=True)
	output_path.write_text(text, encoding="utf-8")
	print(f"Cover letter written to: {output_path}")


def render_text(text: str, output_path: Path) -> None:
	"""Write the cover letter as a plain .txt file (strips markdown bold markers)."""
	output_path.parent.mkdir(parents=True, exist_ok=True)
	clean = text.replace("**", "")
	output_path.write_text(clean, encoding="utf-8")
	print(f"Cover letter written to: {output_path}")


def render_to_stdout(text: str) -> None:
	"""Print the cover letter directly to the terminal."""
	print("\n" + "=" * 60)
	print(text)
	print("=" * 60 + "\n")


def render_docx(
	text: str,
	output_path: Path,
	profile: dict[str, Any],
	job: dict[str, Any],
) -> None:
	"""Render the cover letter as a formatted .docx file.

	Produces a printer-ready document with a coloured header box and a
	two-column body (narrow sidebar on the left, letter text on the right).
	The file can be opened in Word and saved as PDF via File > Save As.
	"""
	output_path.parent.mkdir(parents=True, exist_ok=True)

	contact = profile.get("contact", {})
	name = contact.get("name", "")
	email = contact.get("email", "")
	phone = contact.get("phone", "")
	location = contact.get("location", "")
	company = job.get("company", "")

	# Derive sender credential: prefer education degree, fall back to job title
	education = profile.get("education", [])
	work_history = profile.get("work_history", [])
	if education:
		credential = education[0].get("degree", "")
	elif work_history:
		credential = work_history[0].get("title", "")
	else:
		credential = ""

	doc = Document()

	# Set page margins
	section = doc.sections[0]
	for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
		setattr(section, attr, _MARGIN)

	# --- Header box (1×1 table, full width) ---
	header_table = doc.add_table(rows=1, cols=1)
	header_table.style = "Table Grid"
	_remove_table_borders(header_table)
	hdr_cell = header_table.cell(0, 0)
	_set_cell_background(hdr_cell, _HEADER_FILL)

	p_name = hdr_cell.paragraphs[0]
	p_name.paragraph_format.space_before = Pt(8)
	_add_run(p_name, name, bold=True, size=Pt(16), color=RGBColor(0xFF, 0xFF, 0xFF))

	# Each contact field on its own line
	contact_parts = [x for x in [location, phone, email] if x]
	for i, part in enumerate(contact_parts):
		p_contact = hdr_cell.add_paragraph()
		_add_run(p_contact, part, size=Pt(9), color=RGBColor(0xFF, 0xFF, 0xFF))
		p_contact.paragraph_format.space_after = Pt(8) if i == len(contact_parts) - 1 else Pt(1)

	# Spacer between header and body
	spacer = doc.add_paragraph()
	spacer.paragraph_format.space_after = Pt(4)

	# --- Two-column body table ---
	body_table = doc.add_table(rows=1, cols=2)
	body_table.style = "Table Grid"
	_remove_table_borders(body_table)

	left_cell = body_table.cell(0, 0)
	right_cell = body_table.cell(0, 1)
	# Force fixed table layout so Word respects explicit column widths
	_set_table_fixed_layout(body_table, _SIDEBAR_WIDTH, _CONTENT_WIDTH)
	_set_cell_background(left_cell, _SIDEBAR_FILL)

	# Left sidebar content
	_add_label(left_cell, "RECIPIENT")
	_add_body_text(left_cell, company)
	spacer_p = left_cell.add_paragraph()
	spacer_p.paragraph_format.space_after = Pt(60)
	_add_label(left_cell, "SENDER")
	_add_body_text(left_cell, name, bold=True)
	if credential:
		_add_body_text(left_cell, credential, italic=True, size=Pt(9))

	# Right content: strip markdown bold markers, split into paragraphs
	clean_text = text.replace("**", "")
	paragraphs = [p.strip() for p in clean_text.split("\n\n") if p.strip()]
	first = True
	for para_text in paragraphs:
		if first:
			p = right_cell.paragraphs[0]  # reuse the auto-created empty paragraph
			first = False
		else:
			p = right_cell.add_paragraph()
		_add_run(p, para_text, size=Pt(10))
		p.paragraph_format.space_after = Pt(8)

	# --- Footer ---
	footer = section.footer
	fp = footer.paragraphs[0]
	fp.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
	_add_run(fp, _FOOTER_TEXT, size=Pt(9), color=RGBColor(0x40, 0x40, 0x40))

	try:
		doc.save(output_path)
	except PermissionError:
		raise SystemExit(
			f"Error: cannot write to '{output_path}' — the file is open in another application. Close it and try again."
		)
	print(f"Cover letter written to: {output_path}")


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------


def _set_table_fixed_layout(table: Any, left_width: Any, right_width: Any) -> None:
	"""Force fixed table layout and set explicit column widths (in twips)."""
	# Convert EMU to twips: 1 cm = 567 twips; Cm() stores in EMUs (914400 per inch, 360000 per cm)
	# python-docx Cm() returns EMU; 1 twip = 914400/1440 EMU = 635 EMU
	left_twips = int(left_width / 635)
	right_twips = int(right_width / 635)
	total_twips = left_twips + right_twips

	tbl = table._tbl
	tblPr = tbl.tblPr

	# Set total table width
	tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_twips}" w:type="dxa"/>')
	tblPr.append(tblW)

	# Force fixed layout
	tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
	tblPr.append(tblLayout)

	# Set individual cell widths
	left_cell = table.cell(0, 0)
	right_cell = table.cell(0, 1)
	for cell, twips in ((left_cell, left_twips), (right_cell, right_twips)):
		tcPr = cell._tc.get_or_add_tcPr()
		tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{twips}" w:type="dxa"/>')
		tcPr.append(tcW)


def _set_cell_background(cell: Any, hex_color: str) -> None:
	"""Set the background fill colour of a table cell via direct XML."""
	shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
	cell._tc.get_or_add_tcPr().append(shd)


def _remove_table_borders(table: Any) -> None:
	"""Remove all visible borders from a table via direct XML."""
	tbl_borders = parse_xml(
		f"<w:tblBorders {nsdecls('w')}>"
		'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
		'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
		'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
		'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
		'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
		'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
		"</w:tblBorders>"
	)
	table._tbl.tblPr.append(tbl_borders)


def _add_run(
	paragraph: Any,
	text: str,
	bold: bool = False,
	italic: bool = False,
	size: Any = None,
	color: Any = None,
) -> Any:
	"""Add a formatted run to a paragraph."""
	run = paragraph.add_run(text)
	run.bold = bold
	run.italic = italic
	if size is not None:
		run.font.size = size
	if color is not None:
		run.font.color.rgb = color
	run.font.name = "Calibri"
	return run


def _add_label(cell: Any, label_text: str) -> None:
	"""Add a small all-caps bold label paragraph (RECIPIENT / SENDER style)."""
	p = cell.add_paragraph()
	p.paragraph_format.space_before = Pt(8)
	p.paragraph_format.space_after = Pt(2)
	run = p.add_run(label_text.upper())
	run.bold = True
	run.font.size = Pt(8)
	run.font.name = "Calibri"
	run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)


def _add_body_text(
	cell: Any,
	text: str,
	bold: bool = False,
	italic: bool = False,
	size: Any = None,
) -> None:
	"""Add a regular text paragraph to a table cell."""
	p = cell.add_paragraph()
	p.paragraph_format.space_after = Pt(2)
	run = p.add_run(text)
	run.bold = bold
	run.italic = italic
	run.font.size = size if size is not None else Pt(10)
	run.font.name = "Calibri"
